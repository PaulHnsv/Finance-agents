"""DOCX text extraction using python-docx."""
from pathlib import Path
from docx import Document

def extract_text_from_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

def extract_tables_from_docx(path: Path) -> list[list[list[str]]]:
    doc = Document(str(path))
    result = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        result.append(rows)
    return result
